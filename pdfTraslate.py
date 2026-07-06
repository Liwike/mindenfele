from pdf2docx import Converter
from docx import Document
from deep_translator import GoogleTranslator
from docx2pdf import convert

import sys
import os
import time


def usage():
    print("\nPDF Fordító EN -> HU")
    print("Használat:")
    print(f"  python {os.path.basename(sys.argv[0])} input.pdf")
    print("\nPélda:")
    print(f"  python {os.path.basename(sys.argv[0])} english_book.pdf")
    sys.exit(1)


def translate_text(text):
    try:
        result = GoogleTranslator(
            source="en",
            target="hu"
        ).translate(text)

        if result is None:
            return text

        return str(result)

    except Exception as e:
        print(f"Fordítási hiba: {e}")
        return text


# Paraméter ellenőrzése
if len(sys.argv) != 2:
    usage()

PDF_INPUT = sys.argv[1]

if not os.path.isfile(PDF_INPUT):
    print(f"Hiba: a fájl nem található: {PDF_INPUT}")
    sys.exit(1)

# Fájlnév generálás
base_name = os.path.splitext(PDF_INPUT)[0]

TEMP_DOCX = f"{base_name}_temp.docx"
HU_DOCX = f"{base_name}_HU.docx"
HU_PDF = f"{base_name}_HU.pdf"

print("1. PDF -> DOCX")

try:
    cv = Converter(PDF_INPUT)
    cv.convert(TEMP_DOCX)
    cv.close()
except Exception as e:
    print(f"PDF konvertálási hiba: {e}")
    sys.exit(1)

print("2. DOCX fordítása")

doc = Document(TEMP_DOCX)

# Bekezdések fordítása
for paragraph in doc.paragraphs:
    for run in paragraph.runs:

        text = run.text

        if not text:
            continue

        text = text.strip()

        if not text:
            continue

        print(f"Fordítás: {text[:50]}")

        translated = translate_text(text)

        if translated:
            run.text = translated

        time.sleep(0.2)

# Táblázatok fordítása
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:

                    text = run.text

                    if not text:
                        continue

                    text = text.strip()

                    if not text:
                        continue

                    translated = translate_text(text)

                    if translated:
                        run.text = translated

                    time.sleep(0.2)

doc.save(HU_DOCX)

print("3. DOCX -> PDF")

try:
    convert(HU_DOCX, HU_PDF)
except Exception as e:
    print(f"PDF generálási hiba: {e}")

# Ideiglenes DOCX törlése
if os.path.exists(TEMP_DOCX):
    os.remove(TEMP_DOCX)

print("\nKész.")
print(f"Bemenet : {PDF_INPUT}")
print(f"DOCX    : {HU_DOCX}")
print(f"PDF     : {HU_PDF}")