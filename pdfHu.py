from pdf2docx import Converter
from docx import Document
from docx2pdf import convert

from transformers import MarianMTModel, MarianTokenizer
import torch

import sys
import os
import time

# --------------------------------------------------
# Marian Modell
# Pl. Helsinki-NLP/opus-mt-en-hu letöltve a ./model alá
# --------------------------------------------------

MODEL_PATH = "./models/opus-mt-en-hu"

device = torch.device(
    "cuda" if torch.cuda.is_available() else "cpu"
)

print(f"Használt eszköz: {device}")

if torch.cuda.is_available():
    print(f"GPU: {torch.cuda.get_device_name(0)}")
else:
    print("CUDA nem érhető el, CPU használata.")

tokenizer = MarianTokenizer.from_pretrained(MODEL_PATH)

try:
    model = MarianMTModel.from_pretrained(
        MODEL_PATH,
        torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32
    ).to(device)
except Exception:
    model = MarianMTModel.from_pretrained(
        MODEL_PATH
    ).to(device)

model.eval()


def usage():
    print("\nPDF Fordító EN -> HU (MarianMT)")
    print("Használat:")
    print(f"  python {os.path.basename(sys.argv[0])} input.pdf")
    print("\nPélda:")
    print(f"  python {os.path.basename(sys.argv[0])} english_book.pdf")
    sys.exit(1)


def translate_text(text):

    try:

        inputs = tokenizer(
            text,
            return_tensors="pt",
            truncation=True,
            max_length=512
        )

        inputs = {
            k: v.to(device)
            for k, v in inputs.items()
        }

        with torch.no_grad():

            output = model.generate(
                **inputs,
                max_length=512,
                num_beams=4,
                early_stopping=True
            )

        translated = tokenizer.decode(
            output[0],
            skip_special_tokens=True
        )

        if translated:
            return translated

        return text

    except Exception as e:

        print(f"Fordítási hiba: {e}")

        return text


# --------------------------------------------------
# Paraméter ellenőrzés
# --------------------------------------------------

if len(sys.argv) != 2:
    usage()

PDF_INPUT = sys.argv[1]

if not os.path.isfile(PDF_INPUT):
    print(f"Hiba: a fájl nem található: {PDF_INPUT}")
    sys.exit(1)

# --------------------------------------------------
# Kimeneti fájlnevek
# --------------------------------------------------

base_name = os.path.splitext(PDF_INPUT)[0]

TEMP_DOCX = f"{base_name}_temp.docx"
HU_DOCX = f"{base_name}_HU.docx"
HU_PDF = f"{base_name}_HU.pdf"

# --------------------------------------------------
# PDF -> DOCX
# --------------------------------------------------

print("1. PDF -> DOCX")

try:

    cv = Converter(PDF_INPUT)
    cv.convert(TEMP_DOCX)
    cv.close()

except Exception as e:

    print(f"PDF konvertálási hiba: {e}")
    sys.exit(1)

# --------------------------------------------------
# Fordítás
# --------------------------------------------------

print("2. DOCX fordítása")

doc = Document(TEMP_DOCX)

translated_count = 0

# Bekezdések

for paragraph in doc.paragraphs:

    for run in paragraph.runs:

        text = run.text

        if not text:
            continue

        text = text.strip()

        if not text:
            continue

        print(f"Fordítás: {text[:60]}")

        translated = translate_text(text)

        if translated:
            run.text = translated
            translated_count += 1

# Táblázatok

for table in doc.tables:

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:

                for run in paragraph.runs:

                    text = run.text

                    if not text:
                        continue

                    text = text.strip()

                    if not text:
                        continue

                    translated = translate_text(text)

                    if translated:
                        run.text = translated
                        translated_count += 1

# --------------------------------------------------
# DOCX mentés
# --------------------------------------------------

doc.save(HU_DOCX)

# --------------------------------------------------
# DOCX -> PDF
# --------------------------------------------------

print("3. DOCX -> PDF")

try:

    convert(HU_DOCX, HU_PDF)

except Exception as e:

    print(f"PDF generálási hiba: {e}")

# --------------------------------------------------
# Temp törlése
# --------------------------------------------------

try:

    if os.path.exists(TEMP_DOCX):
        os.remove(TEMP_DOCX)

except Exception:
    pass

# --------------------------------------------------
# Kész
# --------------------------------------------------

print("\nKész.")
print(f"Fordított elemek: {translated_count}")
print(f"Bemenet : {PDF_INPUT}")
print(f"DOCX    : {HU_DOCX}")
print(f"PDF     : {HU_PDF}")
